from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.schemas import ClassifiedOutput, ArticleWithContent
from app.utils.logging import get_logger

logger = get_logger(__name__)


def generate_docx(
    classification: ClassifiedOutput,
    articles: list[ArticleWithContent],
    output_path: Path,
    date_str: str,
) -> Path:
    """Generate the DOCX table of contents matching the required template.

    Template:
    [더벨]
    1. Deal
       A. 경영권 인수 및 매각, 투자 유치
          (1) article title
       B. 투자회수
          (1) article title
       C. 기타
          (1) article title
    2. Industry
       A. E&F 포트폴리오 관련 산업 업계 동향
          - 환경/폐기물
            (1) article title
          - 건설/부동산
            (1) article title
          - 바이오/헬스케어
            (1) article title
       B. 기타 주요 산업 관련 업계 동향
          (1) article title
    3. Fundraising, LP 이슈 및 GP 선정
       (1) article title
    """
    articles_map = {a.info.id: a for a in articles}

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "맑은 고딕"
    font.size = Pt(10)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"(더벨) Daily News Clipping {date_str}")
    run.bold = True
    run.font.size = Pt(16)

    doc.add_paragraph()  # Blank line

    # Helper to add articles under a section
    def add_articles(article_ids: list[str], indent_level: int = 2):
        if not article_ids:
            return
        count = 0
        for aid in article_ids:
            a = articles_map.get(aid)
            if a:
                count += 1
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(indent_level * 1.0)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(f"({count})\t{a.info.title}")
                run.font.size = Pt(10)

    for cat in classification.categories:
        # Main category number
        cat_num = classification.categories.index(cat) + 1

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(f"{cat_num}.\t{cat.name}")
        run.bold = True
        run.font.size = Pt(12)

        # Direct articles under category (e.g., Fundraising)
        if cat.articles:
            add_articles(cat.articles, indent_level=1)

        # Subcategories
        for sub_idx, sub in enumerate(cat.subcategories):
            sub_letter = chr(65 + sub_idx)  # A, B, C...

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.space_before = Pt(6)
            run = p.add_run(f"{sub_letter}.\t{sub.name}")
            run.bold = True
            run.font.size = Pt(11)

            # Direct articles under subcategory
            if sub.articles:
                add_articles(sub.articles, indent_level=2)

            # Sub-items (e.g., 환경/폐기물, 건설/부동산, etc.)
            for sub_item in sub.sub_items:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(2.0)
                p.paragraph_format.space_before = Pt(4)
                run = p.add_run(f"-\t{sub_item.name}")
                run.font.size = Pt(10)

                if sub_item.articles:
                    add_articles(sub_item.articles, indent_level=3)

    doc.save(str(output_path))
    logger.info(f"DOCX generated: {output_path}")
    return output_path
